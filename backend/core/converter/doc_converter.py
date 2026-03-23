"""
DOC/DOCX转换器
"""
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional
from loguru import logger

from .base import BaseConverter, ConversionResult, DocumentType


class DocConverter(BaseConverter):
    """DOC/DOCX转换器"""
    
    def supports(self, doc_type: DocumentType) -> bool:
        return doc_type in [DocumentType.DOC, DocumentType.DOCX]
    
    async def convert(self, file_path: str) -> ConversionResult:
        """
        DOC/DOCX转换流程:
        1. DOC → DOCX (如果需要)
        2. DOCX → Markdown
        """
        try:
            path = Path(file_path)
            
            # 如果是DOC格式，先转换为DOCX
            if path.suffix.lower() == '.doc':
                docx_path = await self._convert_doc_to_docx(file_path)
                if not docx_path:
                    return ConversionResult(
                        success=False,
                        error_message="无法将DOC转换为DOCX"
                    )
            else:
                docx_path = file_path
            
            # DOCX转Markdown
            markdown_content = await self._convert_docx_to_markdown(docx_path)
            
            # 提取元数据
            metadata = {
                'source_type': path.suffix.lower()[1:],
                'title': path.stem,
                'word_count': len(markdown_content.split()),
                'char_count': len(markdown_content)
            }
            
            return ConversionResult(
                success=True,
                content=markdown_content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"DOC/DOCX转换失败: {e}")
            return ConversionResult(
                success=False,
                error_message=str(e)
            )
    
    async def _convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """将DOC转换为DOCX"""
        # 方法1: 使用LibreOffice
        try:
            return await self._convert_with_libreoffice(doc_path)
        except Exception as e:
            logger.warning(f"LibreOffice转换失败: {e}")
        
        # 方法2: 使用antiword
        try:
            return await self._convert_with_antiword(doc_path)
        except Exception as e:
            logger.warning(f"antiword转换失败: {e}")
        
        # 方法3: 直接读取文本（降级方案）
        logger.warning("使用降级方案读取DOC内容")
        return None
    
    async def _convert_with_libreoffice(self, doc_path: str) -> str:
        """使用LibreOffice转换"""
        doc_path = Path(doc_path)
        output_dir = doc_path.parent
        
        # 检查LibreOffice是否可用
        try:
            subprocess.run(['soffice', '--version'], capture_output=True, check=True)
        except (subprocess.CalledProcessError, FileNotFoundError):
            raise Exception("LibreOffice未安装")
        
        # 执行转换
        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(output_dir),
            str(doc_path)
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode != 0:
            raise Exception(f"LibreOffice转换失败: {result.stderr}")
        
        output_path = output_dir / f"{doc_path.stem}.docx"
        if output_path.exists():
            return str(output_path)
        else:
            raise Exception("LibreOffice输出文件不存在")
    
    async def _convert_with_antiword(self, doc_path: str) -> str:
        """使用antiword提取文本"""
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0:
                # 创建临时markdown文件
                with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
                    f.write(result.stdout)
                    return f.name
            else:
                raise Exception(f"antiword失败: {result.stderr}")
        except FileNotFoundError:
            raise Exception("antiword未安装")
    
    async def _convert_docx_to_markdown(self, docx_path: str) -> str:
        """将DOCX转换为Markdown"""
        try:
            # 尝试使用python-docx
            from docx import Document
            
            doc = Document(docx_path)
            markdown_lines = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    markdown_lines.append('')
                    continue
                
                # 根据样式判断标题级别
                try:
                    style_name = para.style.name if para.style else ''
                except:
                    style_name = ''
                
                if 'Heading 1' in style_name or style_name.startswith('标题 1'):
                    markdown_lines.append(f'# {text}')
                elif 'Heading 2' in style_name or style_name.startswith('标题 2'):
                    markdown_lines.append(f'## {text}')
                elif 'Heading 3' in style_name or style_name.startswith('标题 3'):
                    markdown_lines.append(f'### {text}')
                else:
                    markdown_lines.append(text)
            
            # 处理表格
            for table in doc.tables:
                markdown_lines.append('')  # 空行
                markdown_lines.append(self._convert_table_to_markdown(table))
                markdown_lines.append('')
            
            result = '\n'.join(markdown_lines)
            logger.info(f"DOCX转换成功，共 {len(result)} 字符")
            return result
            
        except ImportError:
            logger.warning("python-docx未安装，使用基础文本提取")
            return await self._extract_text_basic(docx_path)
        except Exception as e:
            logger.error(f"DOCX转Markdown失败: {e}")
            return await self._extract_text_basic(docx_path)
    
    def _convert_table_to_markdown(self, table) -> str:
        """将docx表格转换为Markdown表格"""
        lines = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
            
            # 添加表头分隔符
            if i == 0:
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        
        return '\n'.join(lines)
    
    async def _extract_text_basic(self, docx_path: str) -> str:
        """基础文本提取（降级方案）"""
        try:
            # 尝试作为zip读取（DOCX实际上是zip文件）
            import zipfile
            from xml.etree import ElementTree as ET
            
            NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            with zipfile.ZipFile(docx_path) as zf:
                with zf.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    texts = []
                    for elem in root.iter():
                        if elem.tag == f"{{{NS['w']}}}t":
                            if elem.text:
                                texts.append(elem.text)
                    
                    return '\n\n'.join(texts)
                    
        except Exception as e:
            logger.error(f"基础文本提取失败: {e}")
            return f"[无法提取文档内容: {e}]"
